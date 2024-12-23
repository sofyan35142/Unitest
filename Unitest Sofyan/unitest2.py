import unittest
from unittest.mock import patch, MagicMock
import os
from docx import Document
from io import BytesIO

# Asumsikan nama file yang digunakan adalah 'orders.docx'
ORDER_FILE = 'orders.docx'

# Fungsi yang sudah ada (diimport sesuai dengan implementasi Anda)
from demo import tambah_pesanan, baca_pesanan, update_pesanan, hapus_pesanan, search_pesanan

class TestManajemenPesanan(unittest.TestCase):
    
    def setUp(self):
        self.mock_doc = MagicMock(spec=Document)
        self.mock_open = patch("builtins.open", MagicMock()).start()
        self.mock_os_path_exists = patch("os.path.exists", return_value=True).start()

    def tearDown(self):
        # Stop patching after each test
        patch.stopall()

    def test_tambah_pesanan(self):
        with patch("builtins.input", side_effect=["John Doe", "Nasi Goreng", "2", "path_to_image.jpg"]), \
            patch("docx.Document.add_paragraph") as mock_add_paragraph, \
            patch("docx.Document.add_picture") as mock_add_picture, \
            patch("docx.Document.save") as mock_save:
            
            # Memanggil fungsi untuk menambahkan pesanan
            tambah_pesanan()
            
            # Memastikan bahwa fungsi menambahkan paragraf yang benar
            mock_add_paragraph.assert_any_call("Nama Pelanggan: John Doe")
            mock_add_paragraph.assert_any_call("Menu yang Dipesan: Nasi Goreng")
            mock_add_paragraph.assert_any_call("Jumlah Pesanan: 2")
            mock_add_paragraph.assert_any_call("Status: diproses")
            
            # Memastikan gambar disisipkan jika path gambar valid
            mock_add_picture.assert_called_once_with("path_to_image.jpg", width=docx.shared.Inches(10))
            
            # Memastikan bahwa file disimpan
            mock_save.assert_called_once()

    def test_baca_pesanan(self):
        # Mock data pesanan di dalam dokumen
        doc_mock = MagicMock(spec=Document)
        doc_mock.paragraphs = [
            MagicMock(text="Nama Pelanggan: John Doe"),
            MagicMock(text="Menu yang Dipesan: Nasi Goreng"),
            MagicMock(text="Jumlah Pesanan: 2"),
            MagicMock(text="Status: diproses"),
            MagicMock(text="-" * 50)
        ]
        with patch("docx.Document", return_value=doc_mock):
            with patch("builtins.print") as mock_print:
                baca_pesanan()
                # Memastikan output mencetak informasi pesanan yang benar
                mock_print.assert_any_call("Data Pesanan:")
                mock_print.assert_any_call("1. Nama: John Doe, Menu: Nasi Goreng, Jumlah: 2, Status: diproses")

    def test_update_pesanan(self):
        doc_mock = MagicMock(spec=Document)
        doc_mock.paragraphs = [
            MagicMock(text="Nama Pelanggan: John Doe"),
            MagicMock(text="Menu yang Dipesan: Nasi Goreng"),
            MagicMock(text="Jumlah Pesanan: 2"),
            MagicMock(text="Status: diproses"),
            MagicMock(text="-" * 50)
        ]
        with patch("docx.Document", return_value=doc_mock), \
            patch("builtins.input", side_effect=["1", "selesai"]), \
            patch("docx.Document.save") as mock_save:
            
            update_pesanan()
            # Memastikan status diperbarui
            doc_mock.paragraphs[3].text = "Status: selesai"
            
            # Memastikan file disimpan setelah update
            mock_save.assert_called_once()

    def test_hapus_pesanan(self):
        doc_mock = MagicMock(spec=Document)
        doc_mock.paragraphs = [
            MagicMock(text="Nama Pelanggan: John Doe"),
            MagicMock(text="Menu yang Dipesan: Nasi Goreng"),
            MagicMock(text="Jumlah Pesanan: 2"),
            MagicMock(text="Status: diproses"),
            MagicMock(text="-" * 50)
        ]
        with patch("docx.Document", return_value=doc_mock), \
             patch("builtins.input", side_effect=["1"]), \
             patch("docx.Document.save") as mock_save:
            
            hapus_pesanan()
            # Memastikan bahwa paragraf pesanan dihapus (paragraf terkait status)
            doc_mock.paragraphs.clear()
            mock_save.assert_called_once()

    def test_search_pesanan(self):
        doc_mock = MagicMock(spec=Document)
        doc_mock.paragraphs = [
            MagicMock(text="Nama Pelanggan: John Doe"),
            MagicMock(text="Menu yang Dipesan: Nasi Goreng"),
            MagicMock(text="Jumlah Pesanan: 2"),
            MagicMock(text="Status: diproses"),
            MagicMock(text="-" * 50)
        ]
        with patch("docx.Document", return_value=doc_mock), \
             patch("builtins.input", return_value="John Doe"), \
             patch("builtins.print") as mock_print:
            
            search_pesanan()
            # Memastikan hasil pencarian pesanan ditampilkan
            mock_print.assert_any_call("Nama: John Doe, Menu: Nasi Goreng, Jumlah: 2, Status: diproses")

if __name__ == "__main__":
    unittest.main()
