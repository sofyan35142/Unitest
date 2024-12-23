import os
from docx import Document
# from docx.oxml import parse_xml
# from docx.oxml.ns import qn
import docx.shared

ORDER_FILE = 'orders.docx'

def tambah_pesanan():
    nama = input("Masukkan Nama Pelanggan: ")
    menu = input("Masukkan Menu yang Dipesan: ")
    jumlah = int(input("Masukkan Jumlah Pesanan: "))
    image_path = input("Masukkan Path Gambar: ")
    
    if os.path.exists(ORDER_FILE):
        doc = Document(ORDER_FILE)
    else:
        doc = Document()
        doc.add_heading('Manajemen Pesanan Restoran', 0) 
        print("File 'orders.docx' baru saja dibuat.")
    

    doc.add_paragraph(f"Nama Pelanggan: {nama}")
    doc.add_paragraph(f"Menu yang Dipesan: {menu}")
    doc.add_paragraph(f"Jumlah Pesanan: {jumlah}")
    doc.add_paragraph(f"Status: diproses")
    
    if os.path.exists(image_path):
        doc.add_paragraph('Gambar Pesanan:')
        doc.add_picture(image_path, width=docx.shared.Inches(10))  # Menyisipkan gambar dengan lebar 2 inci
    else:
        doc.add_paragraph(f"Path Gambar tidak valid: {image_path}")
    
    doc.add_paragraph('-' * 50)
    
    doc.save(ORDER_FILE)
    
    print("Pesanan berhasil ditambahkan dan disimpan.")

def baca_pesanan():
    if not os.path.exists(ORDER_FILE):
        print("Belum ada pesanan.")
        return
    
    doc = Document(ORDER_FILE)
    
    if len(doc.paragraphs) == 0:
        print("Belum ada pesanan.")
        return
    
    print("Data Pesanan:")
    order_number = 1
    order_details = []  
    
    order = {}
    for para in doc.paragraphs:
        if "Nama Pelanggan" in para.text:
            order['Nama'] = para.text.split(":")[1].strip()
        elif "Menu yang Dipesan" in para.text:
            order['Menu'] = para.text.split(":")[1].strip()
        elif "Jumlah Pesanan" in para.text:
            order['Jumlah'] = para.text.split(":")[1].strip()
        elif "Status" in para.text:
            order['Status'] = para.text.split(":")[1].strip()
            
            # Menyelesaikan satu pesanan setelah menemukan status
            # Menambahkan hanya jika semua data sudah lengkap
            if 'Nama' in order and 'Menu' in order and 'Jumlah' in order and 'Status' in order:
                order_details.append(order.copy())
            order = {}  # Reset untuk pesanan selanjutnya
    
    if not order_details:
        print("Tidak ada pesanan yang tersedia.")
    else:
        for idx, order in enumerate(order_details, start=1):
            print(f"{idx}. Nama: {order['Nama']}, Menu: {order['Menu']}, Jumlah: {order['Jumlah']}, Status: {order['Status']}")

def update_pesanan():
    if not os.path.exists(ORDER_FILE):
        print("Belum ada pesanan untuk diubah.")
        return
    
    doc = Document(ORDER_FILE)
    
    if len(doc.paragraphs) == 0:
        print("Belum ada pesanan untuk diubah.")
        return
    
    print("Daftar Pesanan:")
    order_details = []  
    order = {}
    
    for para in doc.paragraphs:
        if "Nama Pelanggan" in para.text:
            order['Nama'] = para.text.split(":")[1].strip()
        elif "Menu yang Dipesan" in para.text:
            order['Menu'] = para.text.split(":")[1].strip()
        elif "Jumlah Pesanan" in para.text:
            order['Jumlah'] = para.text.split(":")[1].strip()
        elif "Status" in para.text:
            order['Status'] = para.text.split(":")[1].strip()
            
            if 'Nama' in order and 'Menu' in order and 'Jumlah' in order and 'Status' in order:
                order_details.append(order.copy())
            order = {}
    

    for idx, order in enumerate(order_details, start=1):
        print(f"{idx}. Nama: {order['Nama']}, Menu: {order['Menu']}, Jumlah: {order['Jumlah']}, Status: {order['Status']}")
    
    if len(order_details) == 0:
        print("Tidak ada pesanan yang dapat diubah.")
        return
    
    order_id = int(input("Pilih nomor pesanan yang ingin diubah: ")) - 1
    if 0 <= order_id < len(order_details):
        new_status = input("Masukkan status baru (diproses/selesai/batal): ")
        order_to_update = order_details[order_id]
        
        # Menemukan paragraf yang sesuai dengan pesanan yang dipilih
        start_idx = None
        order_idx = 0
        for idx, para in enumerate(doc.paragraphs):
            if para.text == f"Nama Pelanggan: {order_to_update['Nama']}":
                start_idx = idx
                break
        
        if start_idx is not None:
            doc.paragraphs[start_idx + 3].text = f"Status: {new_status}"
            doc.save(ORDER_FILE)
            print("Pesanan berhasil diubah.")
        else:
            print("Pesanan tidak ditemukan.")
    else:
        print("Nomor pesanan tidak valid.")


def hapus_pesanan():
    if not os.path.exists(ORDER_FILE):
        print("Belum ada pesanan untuk dihapus.")
        return

    doc = Document(ORDER_FILE)

    if len(doc.paragraphs) == 0:
        print("Belum ada pesanan untuk dihapus.")
        return

    print("Daftar Pesanan:")
    order_details = []
    order = {}

    for para in doc.paragraphs:
        if "Nama Pelanggan" in para.text:
            order['Nama'] = para.text.split(":")[1].strip()
        elif "Menu yang Dipesan" in para.text:
            order['Menu'] = para.text.split(":")[1].strip()
        elif "Jumlah Pesanan" in para.text:
            order['Jumlah'] = para.text.split(":")[1].strip()
        elif "Status" in para.text:
            order['Status'] = para.text.split(":")[1].strip()
            if 'Nama' in order and 'Menu' in order and 'Jumlah' in order and 'Status' in order:
                order_details.append(order.copy())
            order = {}

    for idx, order in enumerate(order_details, start=1):
        print(f"{idx}. Nama: {order['Nama']}, Menu: {order['Menu']}, Jumlah: {order['Jumlah']}, Status: {order['Status']}")

    if len(order_details) == 0:
        print("Tidak ada pesanan yang dapat dihapus.")
        return

    order_id = int(input("Pilih nomor pesanan yang ingin dihapus: ")) - 1
    if 0 <= order_id < len(order_details):
        order_to_delete = order_details[order_id]

        if order_to_delete['Status'].lower() == 'batal':
            start_idx = None
            end_idx = None

            for idx, para in enumerate(doc.paragraphs):
                if para.text == f"Nama Pelanggan: {order_to_delete['Nama']}":
                    start_idx = idx
                if start_idx is not None and para.text.startswith('-' * 50):
                    end_idx = idx
                    break

            if start_idx is not None and end_idx is not None:
                # Hapus paragraf dalam rentang tertentu
                for idx in range(end_idx, start_idx - 1, -1):
                    p = doc.paragraphs[idx]._element
                    p.getparent().remove(p)

                # Menghapus gambar terkait
                for rel in list(doc.part.rels.values()):
                    if "image" in rel.target_ref:  # Hanya memeriksa relasi gambar
                        rel_paragraphs = [
                            p.text for p in doc.paragraphs[start_idx:end_idx + 1]
                        ]
                        if any(rel.target_ref in text for text in rel_paragraphs):
                            doc.part.drop_rel(rel.rId)

                doc.save(ORDER_FILE)
                print("Pesanan dan gambar terkait berhasil dihapus.")
            else:
                print("Pesanan tidak ditemukan.")
        else:
            print("Pesanan dengan status selain 'batal' tidak bisa dihapus.")
    else:
        print("Nomor pesanan tidak valid.")
def search_pesanan():
    if not os.path.exists(ORDER_FILE):
        print("Belum ada pesanan untuk dicari.")
        return
    
    doc = Document(ORDER_FILE)
    
    if len(doc.paragraphs) == 0:
        print("Belum ada pesanan untuk dicari.")
        return
    
    search_nama = input("Masukkan nama pelanggan yang dicari: ")
    found = False
    
    order_details = {} 
    
    for para in doc.paragraphs:
        if "Nama Pelanggan" in para.text:
            nama = para.text.split(":")[1].strip() 
        if "Menu yang Dipesan" in para.text:
            menu = para.text.split(":")[1].strip() 
        if "Jumlah Pesanan" in para.text:
            jumlah = para.text.split(":")[1].strip()  
        if "Status" in para.text:
            status = para.text.split(":")[1].strip()  
            
            if search_nama.lower() in nama.lower():
                order_details['Nama'] = nama
                order_details['Menu'] = menu
                order_details['Jumlah'] = jumlah
                order_details['Status'] = status
                found = True
    
    if found:
        print(f"Nama: {order_details['Nama']}, Menu: {order_details['Menu']}, Jumlah: {order_details['Jumlah']}, Status: {order_details['Status']}")
    else:
        print(f"Pesanan dengan nama {search_nama} tidak ditemukan.")


def menu():
    while True:
        print("\nManajemen Pesanan Restoran")
        print("1. Tambah Pesanan")
        print("2. Tampilkan Pesanan")
        print("3. Update Pesanan")
        print("4. Hapus Pesanan")
        print("5. Cari Pesanan (Opsional)")
        print("6. Keluar")
        
        choice = input("Pilih menu: ")
        
        if choice == '1':
            tambah_pesanan()
        elif choice == '2':
            baca_pesanan()
        elif choice == '3':
            update_pesanan()
        elif choice == '4':
            hapus_pesanan()
        elif choice == '5':
            search_pesanan()
        elif choice == '6':
            print("Terima kasih telah menggunakan program ini!")
            break
        else:
            print("Pilihan tidak valid. Coba lagi.")

menu()
